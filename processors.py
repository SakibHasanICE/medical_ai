"""File processors"""

import os
import base64
from abc import ABC, abstractmethod
from typing import Optional

import PyPDF2
import docx
from PIL import Image
import pytesseract
from openai import OpenAI

from config import Config

# Set Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = Config.TESSERACT_PATH


class BaseProcessor(ABC):
    """Base processor interface"""
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        pass
    
    @abstractmethod
    def get_page_count(self, file_path: str) -> int:
        pass
    
    @abstractmethod
    def get_file_type(self) -> str:
        pass


class PDFProcessor(BaseProcessor):
    """PDF processor"""
    
    def extract_text(self, file_path: str) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def get_page_count(self, file_path: str) -> int:
        with open(file_path, 'rb') as file:
            return len(PyPDF2.PdfReader(file).pages)
    
    def get_file_type(self) -> str:
        return "pdf"


class DOCXProcessor(BaseProcessor):
    """DOCX processor"""
    
    def extract_text(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        
        return text.strip()
    
    def get_page_count(self, file_path: str) -> int:
        doc = docx.Document(file_path)
        return max(1, len(doc.paragraphs) // 30)
    
    def get_file_type(self) -> str:
        return "docx"


class ImageProcessor(BaseProcessor):
    """Image processor with OCR - Tesseract first, then Vision API"""
    
    def __init__(self, openai_client: Optional[OpenAI] = None):
        self.openai_client = openai_client
    
    def extract_text(self, file_path: str) -> str:
        """Extract text - tries Tesseract first, then Vision API as fallback"""
        # Try Tesseract first with French + English support
        try:
            tesseract_text = self._extract_with_tesseract(file_path)
            
            # If Tesseract gives good results (more than 50 characters), use it
            if len(tesseract_text.strip()) > 50:
                print("✓ Text extracted using Tesseract OCR (French + English)")
                return tesseract_text
            else:
                print("⚠ Tesseract result too short, trying Vision API...")
                raise ValueError("Tesseract extraction insufficient")
        except Exception as e:
            print(f"⚠ Tesseract failed: {str(e)}, trying Vision API...")
        
        # Fallback to Vision API if Tesseract fails or gives poor results
        if self.openai_client:
            try:
                vision_text = self._extract_with_vision(file_path)
                print("✓ Text extracted using OpenAI Vision API")
                return vision_text
            except Exception as e:
                print(f"✗ Vision API failed: {str(e)}")
                return ""
        
        return ""
    
    def _extract_with_tesseract(self, file_path: str) -> str:
        """Extract using Tesseract OCR with French + English support"""
        image = Image.open(file_path)
        
        # Try with French + English languages
        try:
            # Use configured languages from Config
            text = pytesseract.image_to_string(image, lang=Config.TESSERACT_LANGUAGES)
            
            # If result is too short, try with English only as fallback
            if len(text.strip()) < 20:
                print("⚠ French+English failed, trying English only...")
                text = pytesseract.image_to_string(image, lang='eng')
        except Exception as e:
            print(f"⚠ Tesseract language error: {e}, trying default...")
            # Fallback to default language
            text = pytesseract.image_to_string(image)
        
        return text.strip()
    
    def _extract_with_vision(self, file_path: str) -> str:
        """Extract using Vision API with multilingual support"""
        with open(file_path, 'rb') as f:
            base64_image = base64.b64encode(f.read()).decode('utf-8')
        
        response = self.openai_client.chat.completions.create(
            model=Config.VISION_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text", 
                        "text": "Extract all text from this image. Preserve the original language (English or French). Return only the extracted text."
                    },
                    {
                        "type": "image_url", 
                        "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                    }
                ]
            }],
            max_tokens=4096
        )
        return response.choices[0].message.content.strip()
    
    def get_page_count(self, file_path: str) -> int:
        return 1
    
    def get_file_type(self) -> str:
        return "image"


class ProcessorFactory:
    """Factory to create processors"""
    
    @staticmethod
    def create(file_path: str, openai_client: Optional[OpenAI] = None) -> BaseProcessor:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return PDFProcessor()
        elif ext in ('.docx', '.doc'):
            return DOCXProcessor()
        elif ext in ('.jpg', '.jpeg', '.png'):
            return ImageProcessor(openai_client)
        else:
            raise ValueError(f"Unsupported file type: {ext}")